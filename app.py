# Data Clean, Axis to Plots Visualization UI

import pandas as pd
import streamlit as st
import plost 

st.set_page_config(layout='wide', page_title='Data Analyzer', page_icon=':bar_chart:')

with st.sidebar: 
    st.title("Data Analyzer")
    dataset = st.file_uploader('Upload Your Dataset',type=['csv','xlsx','xls'])

#@st.cache_data
def load_data(dataset):
    # file_extension = os.path.splitext(dataset)[-1]
    file_extension = dataset.type.split('/')[1]

    if file_extension == "csv":
        data = pd.read_csv(dataset)
    # elif file_extension == "xlsx":
    elif file_extension == "vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        data = pd.read_excel(dataset)
    # elif file_extension == "xls":
    elif file_extension == "vnd.ms-excel":
        data = pd.read_excel(dataset)
    else:
        st.info('Upload csv, xls or xlsx file')

    return data

if dataset is None:
    st.stop()

dataset_text = dataset.name.split(".")[0] 
cleaned_text = dataset_text.replace('-', ' ').replace(',', '').replace('.', '').replace('_', ' ')
dataset_name = ' '.join(word.capitalize() for word in cleaned_text.split())

df = load_data(dataset)

# Remove Null Columns and Rows
df = df.dropna(axis=1, thresh= 0.7*len(df))  
df = df.dropna()    

with st.sidebar:
    ''
    nrows = st.slider('No.of Rows',10,len(df))

df = df.head(nrows)
column_names = df.columns.tolist()

int_columns = []
str_columns = []
bool_columns = []

for column_name, dtype in df.dtypes.items():
    if dtype == 'int64':
        int_columns.append(column_name)
    elif dtype == 'float64':
        int_columns.append(column_name)
    elif dtype == 'object':
        str_columns.append(column_name)
    elif dtype == 'bool':
        bool_columns.append(column_name)

st.header('Data Analyzer')

tab1, tab2 = st.tabs(['Dataset Preview','Plots Visualization'])

with tab1:
    st.write(df.head(7))
    st.write(df.shape)

with tab2:
    col1, col2, col10, col20 = st.columns([3,3,1,3])

    with col1:
        xaxis = st.selectbox("X-axis", column_names)
    with col2:
        yaxis = st.selectbox("Y-axis", column_names)
    with col20:
        ''
        st.subheader('Customize')

    col3, col4 = st.columns([5,2])
            
    with col4:

        plots = ['Line Chart', 'Histogram', 'Scatter Plot', 'Bar Chart', 'Pie Chart']

        if xaxis and yaxis in int_columns:
            plots = ['Histogram', 'Line Chart', 'Scatter Plot']
        
        elif xaxis and yaxis in str_columns:
            plots = ['Pie Chart','Bar Chart']

        elif xaxis in int_columns and yaxis in str_columns:
            plots = ['Bar Chart']
        
        elif xaxis in str_columns and yaxis in int_columns:
            plots = ['Bar Chart']

        elif xaxis in int_columns and yaxis == 'None':
            plots = ['Histogram']
        
        plot = st.selectbox("Plot", plots)

        if plot == 'Line Chart':
            num = st.slider("Opacity", 0.0, -1.0)
        elif plot == 'Histogram':
            sel = st.selectbox('Type', ['None','count', 'distinct', 'sum', 'mean', 'median', 'max', 'min', 'valid', 'missing'])
            ''
        elif plot == 'Scatter Plot':
            siz = st.slider('Size', 0, 100)
        elif plot == 'Pie Chart':
            xaxis = st.selectbox("X-axis", str_columns)
            colraxis = st.selectbox("Color", column_names)
            yaxis = None
        elif plot == 'Bar Chart':
            colraxis = st.selectbox("Color", column_names)


        col41, col42 = st.columns(2)
        
        with col42:
            if plot == 'Line Chart':
                color = st.checkbox("Color")
                with col41:
                    if color:
                        colr = st.color_picker("Pick a color")
                    else: 
                        colr = None
            elif plot == 'Histogram':
                with col41:
                    ycheck = st.checkbox('Y-axis')
                    if not ycheck:
                        yaxis = None
                        sel = 'count'

            rep = st.button("Download")
        
        
    if yaxis != None:
        titl = plot+' of "'+xaxis+'" and "'+yaxis+'"'
    else:
        titl = plot+' of "'+xaxis+'"'

    with col3:
        ""
        ""
        if plot == 'Line Chart':
            plost.line_chart(
                    data = df,
                    x = xaxis,
                    y = yaxis,
                    color = colr,
                    opacity = num+1,
                    title = titl,
                    legend='bottom')
            
        elif plot == 'Histogram':
            plost.hist(
                    data = df, 
                    x = xaxis,
                    y = yaxis, 
                    aggregate = sel, 
                    bin=None, 
                    title = titl, 
                    legend = 'bottom')
            
        elif plot == 'Scatter Plot':
            plost.scatter_chart(
                    data = df, 
                    x = xaxis, 
                    y = yaxis, 
                    color=None, 
                    size = siz+35, 
                    title = titl, 
                    legend='right')
            
        elif plot == 'Bar Chart':
            plost.bar_chart(
                    data = df,
                    bar = xaxis, 
                    value = yaxis, 
                    color=colraxis, 
                    opacity=None, 
                    group=None, 
                    direction = 'vertical', 
                    title = titl,
                    legend = 'right')
            
        elif plot == 'Pie Chart':
            plost.pie_chart(
                    data = df, 
                    theta = xaxis, 
                    color = colraxis,
                    title = titl, 
                    legend='right')

# Report
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK 
from docx.shared import Pt 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import subprocess

path = 'C:/Users/Gopi/Desktop/'

doc = Document()

doc.sections[0].page_width = Pt(595.276)
doc.sections[0].page_height = Pt(841.890)

font = 'Open Sans' 

def firstpage(df2, file_name):
    heading = doc.add_heading(f"{file_name} Data Report", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = heading.runs[0]
    run1.font.name = font
    run1.font.size = Pt(18)

    # Dataset Overview
    subheading1 = doc.add_heading("Dataset Overview", level=2)
    run2 = subheading1.runs[0]
    run2.font.name = font
    run2.font.size = Pt(14)
    content1 = doc.add_paragraph(f"\nThe given {file_name} dataset is examined and varying information about their features are extracted and visualized in this comprehensive report. This dataset contains {len(df2)} Rows and {len(df.columns)} Columns, where the columns are named as {column_names[0]}, {column_names[1]}, and more. The visualizations done ranges from Line Chart, Histogram, Scatter Plot, Bar Chart and Pie Chart along with their insights. They are detailed further and elaborated.")
    run4 = content1.runs[0]
    run4.font.name = font
    run4.font.size = Pt(12)

    content2 = doc.add_paragraph(f"Embarking on a visual journey, this report delves into insightful visualizations, translating complex data into meaningful narratives. From dynamic dashboards to static representations, each visualization unlocks patterns, empowering decision-makers with nuanced understanding. Join us in exploring the data's story, where each graph and chart reveals a new perspective.")
    run5 = content2.runs[0]
    run5.font.name = font
    run5.font.size = Pt(12)

    content3 = doc.add_paragraph(f"We've employed a mix of visualizations to bring our data to life. Picture the histogram as a detailed snapshot of data patterns, the line chart unveiling trends over time. Bar plots break down categories, pie charts show proportions, and scatterplots reveal connections. These visual aids act as guides, helping us navigate the complexities of our dataset and make more informed decisions.")
    run6 = content3.runs[0]
    run6.font.name = font
    run6.font.size = Pt(12)
    doc.add_paragraph() 



# Line Chart
def add_linechart(data, x, y, hue=None, title="Line Chart", xlabel=None, ylabel=None):

    try:
        plt.figure(figsize=(8, 6))
        sns.lineplot(data=data, x=x, y=y, hue=hue)
        plt.title(title)
        plt.xlabel(xlabel if xlabel else x)
        plt.ylabel(ylabel if ylabel else y)
        
        plt.savefig(path + "linechart.png")
        plt.close()
        
        doc.add_heading(title, level=1)
        doc.add_picture(path + "linechart.png", width=Inches(5.5), height=Inches(4))

        linechart1 = doc.add_paragraph(f"This Line Chart represents the relationship between the X and Y variables over a continuous range. It is useful for showing trends and patterns over time or other continuous intervals. This Chart plots '{x}' on X Axis and '{y}' on Y Axis in a Linear Fashion.")
        run4 = linechart1.runs[0]
        run4.font.name = font
        run4.font.size = Pt(12)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    except Exception as e:
        pass


# Histogram
def add_histogram(data, x, y=None, title="Histogram", xlabel=None, ylabel=None, bins=10):

    try:
        plt.figure(figsize=(8, 6))  

        sns.histplot(data=data, x=x, bins=bins, edgecolor='black')
        plt.xlabel(xlabel if xlabel else x)
        plt.ylabel('Frequency')
        plt.title(title)

        plt.savefig(path + "histogram.png")
        plt.close()

        doc.add_heading(title, level=1)
        doc.add_picture(path + "histogram.png", width=Inches(5.5), height=Inches(4))

        histogram = doc.add_paragraph(f"This Histogram displays the distribution of a single variable. It is useful for understanding the underlying frequency distribution of the data. This Histogram plots '{x}' on X Axis and 'Frequency' on Y Axis.")
        run4 = histogram.runs[0]
        run4.font.name = font
        run4.font.size = Pt(12)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    except Exception as e:
        pass


# Scatter Plot
def add_scatterplot(data, x, y, hue=None, title="Scatter Plot", xlabel=None, ylabel=None):

    try:
        plt.figure(figsize=(8, 6))
        sns.scatterplot(data=data, x=x, y=y, hue=hue)
        plt.title(title)
        plt.xlabel(xlabel if xlabel else x)
        plt.ylabel(ylabel if ylabel else y)
        
        plt.savefig(path + "scatterplot.png")
        plt.close()
        
        doc.add_heading(title, level=1)
        doc.add_picture(path + "scatterplot.png", width=Inches(5.5), height=Inches(4))

        scatterplot = doc.add_paragraph(f"The Scatter Plot visualizes the relationship between two variables. Each point represents an observation in the dataset, making it useful for identifying patterns and outliers. This Scatter Plot plots '{x}' on X Axis and '{y}' on Y Axis.")
        run4 = scatterplot.runs[0]
        run4.font.name = font
        run4.font.size = Pt(12)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    except Exception as e:
        pass


# Bar Chart
def add_barchart(data, x, y, hue=None, title="Bar Chart", xlabel=None, ylabel=None):

    try:
        plt.figure(figsize=(8, 6))
        sns.barplot(data=data, x=x, y=y, hue=hue)
        plt.title(title)
        plt.xlabel(xlabel if xlabel else x)
        plt.ylabel(ylabel if ylabel else y)
        
        plt.savefig(path+"barchart.png")
        plt.close()
        
        doc.add_heading(title, level=1)
        doc.add_picture(path+"barchart.png", width=Inches(5.5), height=Inches(4))

        barchart = doc.add_paragraph(f"The Bar Chart represents the relationship between a categorical variable on the X-axis and a numeric variable on the Y-axis. It is useful for comparing values across different categories. This Bar Chart plots '{x}' on X Axis and '{y}' on Y Axis.")
        run4 = barchart.runs[0]
        run4.font.name = font
        run4.font.size = Pt(12)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    except Exception as e:
        pass


# Pie Chart
def add_piechart(data, labels, title="Pie Chart"):

    try:
        plt.figure(figsize=(8, 6))
        plt.pie(data, labels=labels, autopct='%1.1f%%')
        plt.title(title)
        
        plt.savefig(path + "piechart.png")
        plt.close()
        
        doc.add_heading(title, level=1)
        doc.add_picture(path + "piechart.png", width=Inches(5.5), height=Inches(4))

        piechart = doc.add_paragraph(f"The Pie chart is useful for showing the proportional distribution of a single categorical variable. Each slice represents a category, and the size of each slice corresponds to the proportion of that category in the dataset. This Pie Chart plots the '{str_columns[0]}' in the {dataset_name} Dataset.")
        run4 = piechart.runs[0]
        run4.font.name = font
        run4.font.size = Pt(12)
    except Exception as e:
        pass


# Line Chart and Scatter Plot Axis
if len(int_columns)>1:
    x1 = int_columns[0]
    y1 = int_columns[1]
elif len(int_columns)==1:
    x1 = int_columns[0]
    y1 = int_columns[0]
elif len(int_columns)==0:
    x1 = None
    y1 = None

# Histogram Axis
if len(int_columns)>1:
    x3 = int_columns[0]
    y3 = int_columns[1]
elif len(int_columns)==1:
    x3 = int_columns[0]
    y3 = None
elif len(int_columns)==0:
    x3 = None
    y3 = None

# Bar Chart Axis
if len(str_columns)>1:
    x4 = str_columns[0]
    y4 = str_columns[1]
elif len(str_columns)==1:
    x4 = str_columns[0]
    if len(int_columns)>=1:
        y4 = int_columns[0]
elif len(str_columns)==0:
    if len(int_columns)>1:
        x4 = int_columns[0]
        y4 = int_columns[1]
    else:
        x4 = None
        y4 = None

# Pie Chart Axis
if len(str_columns)>=1:
    x2 = str_columns[0]
elif len(str_columns)==0:
    x2 = None


if rep:
    firstpage(df,dataset_name)

    if x1 != None and y1 != None:
        titl = 'Line Chart'+' of "'+x1+'" and "'+y1+'"'
        add_linechart(df, x=x1, y=y1, title=titl)
    
    if x3 != None and y3 != None:
        titl = 'Histogram'+' of "'+x3+'" and "'+y3+'"'
        add_histogram(df, x=x3, y=y3, title=titl)
    elif x3 != None and y3 == None:
        titl = plot+' of "'+x3+'"'
        add_histogram(df, x=x3, title=titl)

    if x1 != None and y1 != None:
        titl = 'Scatter Plot'+' of "'+x1+'" and "'+y1+'"'
        add_scatterplot(df, x=x1, y=y1, title=titl)

    if x4 != None and y4 != None:
        titl = 'Bar Chart'+' of "'+x4+'" and "'+y4+'"'
        add_barchart(df, x=x4, y=y4, title=titl)
    
    if x2 != None:
        titl = 'Pie Chart'+' of "'+x2+'"'
        d = df[x2].value_counts()
        col = titl.split(' ')[-1]
        add_piechart(d, d.index, title=titl)
        
    doc.save(path+"output.docx")
    plotpng = ['linechart.png','histogram.png','scatterplot.png','barchart.png','piechart.png']
    for i in plotpng:
        try:
            os.remove(path+i)
        except Exception:
            pass
    
    subprocess.Popen(['start', 'winword', path+"output.docx"], shell=True)
